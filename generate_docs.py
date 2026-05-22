import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors

# Define content data for the form
PROJECT_NAME = "Нейротуннель «Туннель Состояний» / Neurotunnel: The State Tunnel"
STAGE = "Концепция"
DIRECTION = "Пространственные и объектные индустрии (а также Цифровые и интерактивные, Звуковые и Музыкальные)"
LEGAL_STATUS = "ИП (ИП Латыпов В. М.) или Физическое лицо"

SHORT_DESC = (
    "Нейротуннель «Туннель Состояний» — интерактивный шлюз на базе Эмотех, "
    "бесконтактно считывающий биометрию человека и гармонизирующий его состояние "
    "генеративным звуком и цветом."
)  # 175 characters (strict limit is 200)

RELEVANCE = (
    "Жители мегаполисов находятся в состоянии хронического информационного стресса и ментального шума. "
    "Приходя в музеи или галереи, люди часто не способны глубоко воспринимать культуру из-за накопленной тревожности. "
    "Наш проект решает эту проблему на стыке эмоциональной архитектуры (Эмотех) и бесконтактной биометрии. "
    "Нейротуннель бережно считывает состояние посетителя и моментально погружает его в компенсирующую звуковую и "
    "цветовую среду, повышая восприимчивость к искусству."
)  # 466 characters (strict limit is 1000)

CURRENT_STATUS = (
    "Разработана архитектурная и технологическая концепция в 5 масштабах (от выставочного модуля до городского моста). "
    "Создан рабочий интерактивный веб-симулятор инсталляции с генеративным звуковым и визуальным движком. "
    "Проведены предварительные исследования технологии бесконтактного rPPG-мониторинга. Имеется концептуальное видение "
    "интеграции проекта на ГЭС-2 и мосту к Красному Октябрю. Подготовлен 3D-видеоролик (walkthrough) проекта."
)

RESOURCES_SOUGHT = [
    "Партнёры и площадки (для пилотирования на ГЭС-2 и мостах Москвы)",
    "Экспертиза и менторство (в части доработки алгоритмов биометрии)",
    "Продюсерское сопровождение (для выхода на крупные городские и частные заказы)"
]

TEAM = [
    {
        "role": "Главный архитектор (Конструктив и Пространство)",
        "name": "Алёна Левицкая",
        "bio": "Руководитель архитектурного бюро «СТРУКТУРА» (с 2015 г.), член Союза архитекторов Москвы с 2009 г. Лицензирована Минкультуры РФ на работу с объектами культурного наследия. Автор концепций пространственной интеграции медиа-арта, куратор образовательных программ в МАРХИ."
    },
    {
        "role": "Технологический директор / Физик (Сенсоры, Оптика и Звук)",
        "name": "Валерий Латыпов",
        "bio": "Магистр наук по физике конденсированного состояния (НИЯУ МИФИ). Специалист в области квантовой оптики, акустики и сенсорных интерфейсов. Разработчик систем бесконтактного мониторинга физиологических показателей на основе rPPG (фотоплетизмографии) и компьютерного зрения."
    }
]

TECH_SPECS = [
    ("Конструктив", "Облегченный алюминиевый модульный каркас быстрого монтажа (время сборки — 4 часа)."),
    ("Система отображения", "Гибкие бесшовные LED-экраны высокого разрешения (шаг пикселя 1.8-2.5 мм) на полу, потолке и стенах."),
    ("Сенсорный массив", "ИК-камеры, RGB-камеры, лазерные сенсоры LiDAR (высокоточное определение координат X/Y/Z)."),
    ("Звуковая система", "Ультранаправленные звуковые прожекторы Audio Spotlight для изоляции звуковых зон."),
    ("Программное обеспечение", "TouchDesigner / Unreal Engine 5 (генеративная графика), WebAudio API / MaxMSP (генерация бинауральных ритмов 6Гц-12Гц на несущей частоте 150Гц).")
]

LINKS = [
    ("Сайт-презентация проекта (Demo-лендинг)", "https://nejrofest-tunnel.vercel.app"),
    ("Видео-презентация проекта (3D Walkthrough)", "https://nejrofest-tunnel.vercel.app/images/walkthrough.mp4"),
    ("Схема оборудования и чертеж (blueprint)", "https://nejrofest-tunnel.vercel.app/images/blueprint.png"),
    ("Вид сбоку (архитектурная интеграция)", "https://nejrofest-tunnel.vercel.app/images/side_view.png"),
    ("Вид изнутри (световая среда)", "https://nejrofest-tunnel.vercel.app/images/pov_inside.png")
]


def create_docx(filename):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Title
    title = doc.add_paragraph()
    run_title = title.add_run(PROJECT_NAME.upper())
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    run_sub = subtitle.add_run("Материалы для заполнения анкеты проекта #Нейрофест2026")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(212, 175, 55) # Gold
    subtitle.paragraph_format.space_after = Pt(18)

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    p_div_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '12')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'CCCCCC')
    p_div_border.append(bottom_border)
    p_div._p.get_or_add_pPr().append(p_div_border)

    def add_section_heading(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        return h

    # Loop through form fields
    fields = [
        ("Наименование проекта *", PROJECT_NAME),
        ("Стадия проекта *", STAGE),
        ("Направление проекта *", DIRECTION),
        ("Юридический статус *", LEGAL_STATUS),
        ("Краткое описание * (лимит 200 символов)", SHORT_DESC),
        ("Актуальность проекта * (лимит 1000 символов)", RELEVANCE),
        ("Статус реализации на сегодняшний день *", CURRENT_STATUS),
    ]

    add_section_heading("ПОЛЯ ДЛЯ ЗАПОЛНЕНИЯ В АНКЕТЕ")

    for field_name, value in fields:
        p_name = doc.add_paragraph()
        run_fn = p_name.add_run(f"■ {field_name}:")
        run_fn.bold = True
        run_fn.font.color.rgb = RGBColor(71, 85, 105)
        p_name.paragraph_format.space_before = Pt(6)
        p_name.paragraph_format.space_after = Pt(2)
        p_name.paragraph_format.keep_with_next = True

        p_val = doc.add_paragraph()
        p_val.paragraph_format.left_indent = Inches(0.2)
        run_val = p_val.add_run(value)
        run_val.font.color.rgb = RGBColor(15, 23, 42)
        p_val.paragraph_format.space_after = Pt(10)

    # Resources
    p_res_title = doc.add_paragraph()
    run_rt = p_res_title.add_run("■ Какой основной ресурс вы ищете в Акселераторе? * (отметить галочками)")
    run_rt.bold = True
    run_rt.font.color.rgb = RGBColor(71, 85, 105)
    p_res_title.paragraph_format.space_before = Pt(6)
    p_res_title.paragraph_format.space_after = Pt(2)
    p_res_title.paragraph_format.keep_with_next = True

    for res in RESOURCES_SOUGHT:
        p_res = doc.add_paragraph()
        p_res.paragraph_format.left_indent = Inches(0.2)
        p_res.paragraph_format.space_after = Pt(3)
        run_chk = p_res.add_run("✓  ")
        run_chk.bold = True
        run_chk.font.color.rgb = RGBColor(212, 175, 55)
        p_res.add_run(res)

    p_div2 = doc.add_paragraph()
    p_div2.paragraph_format.space_before = Pt(14)
    p_div2.paragraph_format.space_after = Pt(14)
    p_div2_border = OxmlElement('w:pBdr')
    bottom_border2 = OxmlElement('w:bottom')
    bottom_border2.set(qn('w:val'), 'single')
    bottom_border2.set(qn('w:sz'), '4')
    bottom_border2.set(qn('w:space'), '1')
    bottom_border2.set(qn('w:color'), 'E2E8F0')
    p_div2_border.append(bottom_border2)
    p_div2._p.get_or_add_pPr().append(p_div2_border)

    # 4. Tech Specs Table
    add_section_heading("ДОПОЛНИТЕЛЬНО: ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Компонент'
    hdr_cells[1].text = 'Техническое описание'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(15, 23, 42)
        cell.width = Inches(2.2) if cell == hdr_cells[0] else Inches(4.8)

    for item, desc in TECH_SPECS:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = desc
        row_cells[0].width = Inches(2.2)
        row_cells[1].width = Inches(4.8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Team
    add_section_heading("КОМАНДА ПРОЕКТА")
    for t in TEAM:
        p = doc.add_paragraph()
        run_name = p.add_run(f"{t['name']} — {t['role']}\n")
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(15, 23, 42)
        run_bio = p.add_run(t['bio'])
        run_bio.font.size = Pt(10)
        run_bio.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_after = Pt(10)

    # 6. Links
    add_section_heading("МАТЕРИАЛЫ И ССЫЛКИ ДЛЯ ПРИКРЕПЛЕНИЯ К ЗАЯВКЕ")
    for name, url in LINKS:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        run_bullet = p.add_run("•  ")
        run_bullet.bold = True
        run_bullet.font.color.rgb = RGBColor(212, 175, 55)
        run_name = p.add_run(f"{name}: ")
        run_name.bold = True
        run_url = p.add_run(url)
        run_url.font.color.rgb = RGBColor(14, 165, 233)
        run_url.underline = True

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_foot.add_run("Документ подготовлен для Алёны Левицкой и Валерия Латыпова. Скопируйте нужные разделы напрямую в веб-форму.")
    run_foot.font.size = Pt(8.5)
    run_foot.font.italic = True
    run_foot.font.color.rgb = RGBColor(148, 163, 184)

    doc.save(filename)
    print(f"Word document saved to {filename}")


def create_pdf(filename):
    # Setup document
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    # Font registrations
    pdfmetrics.registerFont(TTFont('Arial', '/System/Library/Fonts/Supplemental/Arial.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', '/System/Library/Fonts/Supplemental/Arial Bold.ttf'))

    # Stylesheet
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=15,
        leading=19,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=4
    )
    
    subtitle_style = ParagraphStyle(
        'DocSub',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#d4af37'),
        spaceAfter=12
    )
    
    section_title_style = ParagraphStyle(
        'DocSecTitle',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#0f172a'),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )

    field_title_style = ParagraphStyle(
        'DocFieldTitle',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=9.5,
        leading=12,
        textColor=colors.HexColor('#475569'),
        spaceBefore=8,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=6,
        leftIndent=10
    )

    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#1e293b'),
        leftIndent=15,
        spaceAfter=3
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#0f172a')
    )

    table_cell_bold = ParagraphStyle(
        'TableCellBold',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor('#334155')
    )

    table_cell_normal = ParagraphStyle(
        'TableCellNormal',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=8,
        leading=10.5,
        textColor=colors.HexColor('#475569')
    )

    link_label_style = ParagraphStyle(
        'LinkLabel',
        parent=styles['Normal'],
        fontName='Arial-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#334155'),
        leftIndent=10
    )

    link_url_style = ParagraphStyle(
        'LinkURL',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor('#0ea5e9'),
        leftIndent=10,
        spaceAfter=4
    )

    story = []

    # Title & Header
    story.append(Paragraph(PROJECT_NAME.upper(), title_style))
    story.append(Paragraph("Шпаргалка для заполнения анкеты проекта #Нейрофест2026", subtitle_style))
    
    # Custom Divider Line
    divider = Table([['']], colWidths=[515], rowHeights=[1.5])
    divider.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#CCCCCC')),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(divider)
    story.append(Spacer(1, 10))

    # Form Fields Section
    story.append(Paragraph("ПОЛЯ ДЛЯ ЗАПОЛНЕНИЯ В ВЕБ-ФОРМЕ", section_title_style))

    fields = [
        ("Наименование проекта *", PROJECT_NAME),
        ("Стадия проекта *", STAGE),
        ("Направление проекта *", DIRECTION),
        ("Юридический статус *", LEGAL_STATUS),
        ("Краткое описание * (лимит 200 символов)", SHORT_DESC),
        ("Актуальность проекта * (лимит 1000 символов)", RELEVANCE),
        ("Статус реализации на сегодняшний день *", CURRENT_STATUS),
    ]

    for field_name, value in fields:
        story.append(Paragraph(f"■ {field_name}", field_title_style))
        story.append(Paragraph(value, body_style))

    story.append(Paragraph("■ Какой основной ресурс вы ищете в Акселераторе? * (отметить галочками)", field_title_style))
    for res in RESOURCES_SOUGHT:
        story.append(Paragraph(f"<font color='#d4af37'><b>✓</b></font>  {res}", bullet_style))

    story.append(Spacer(1, 10))
    story.append(divider)

    # Tech Specs Table
    story.append(Paragraph("ДОПОЛНИТЕЛЬНО: ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ", section_title_style))
    
    table_data = [
        [Paragraph("Компонент", table_header_style), Paragraph("Техническое описание", table_header_style)]
    ]
    for item, desc in TECH_SPECS:
        table_data.append([
            Paragraph(item, table_cell_bold),
            Paragraph(desc, table_cell_normal)
        ])
    
    tech_table = Table(table_data, colWidths=[120, 395])
    tech_table.setStyle(TableStyle([
        ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(tech_table)
    story.append(Spacer(1, 10))

    # Team
    story.append(Paragraph("КОМАНДА ПРОЕКТА", section_title_style))
    for t in TEAM:
        team_text = f"<b>{t['name']}</b> — {t['role']}<br/><font color='#475569' size='8.5'>{t['bio']}</font>"
        story.append(Paragraph(team_text, body_style))

    # Links
    story.append(Paragraph("МАТЕРИАЛЫ И ССЫЛКИ ДЛЯ ПРИКРЕПЛЕНИЯ", section_title_style))
    for name, url in LINKS:
        story.append(Paragraph(f"•  <b>{name}</b>", link_label_style))
        story.append(Paragraph(f"<a href='{url}'>{url}</a>", link_url_style))

    # Build PDF
    doc.build(story)
    print(f"PDF document saved to {filename}")


if __name__ == "__main__":
    output_docx = "/Volumes/Genius Art/Antigravity/nejrofest-tunnel/nejrotonnel_zayavka.docx"
    output_pdf = "/Volumes/Genius Art/Antigravity/nejrofest-tunnel/nejrotonnel_zayavka.pdf"
    
    create_docx(output_docx)
    create_pdf(output_pdf)
